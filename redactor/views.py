from django.shortcuts import render

# Create your views here.
import logging
import tempfile
import traceback
from pathlib import Path
from uuid import uuid4

from django.http import FileResponse, HttpResponse
from django.shortcuts import redirect, render

from docx import Document

logger = logging.getLogger(__name__)


def _import_redactor():
    """Import redaction functions, returning (redact_document, blackout_fn) or raising."""
    from redact_pii import redact_document, blackout_identity_images_in_docx
    return redact_document, blackout_identity_images_in_docx


def index(request):
    if request.method == "POST" and request.FILES.get("document"):
        uploaded = request.FILES["document"]
        step = "starting"

        try:
            step = "importing redact_pii"
            redact_document, blackout_identity_images_in_docx = _import_redactor()

            step = "writing upload to temp file"
            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=".docx"
            ) as temp_input:
                for chunk in uploaded.chunks():
                    temp_input.write(chunk)
                input_path = temp_input.name

            logger.info("Temp input written: %s", input_path)

            output_path = str(
                Path(tempfile.gettempdir()) / f"redacted-{uuid4()}.docx"
            )

            step = "opening document with python-docx"
            doc = Document(input_path)
            logger.info("Document opened successfully")

            step = "running redact_document (NLP + substitution)"
            redacted = redact_document(doc)
            logger.info("Redaction complete")

            step = "saving redacted document"
            redacted.save(output_path)
            logger.info("Saved to: %s", output_path)

            step = "blacking out identity images"
            blackout_identity_images_in_docx(output_path)
            logger.info("Image blackout complete")

            step = "streaming FileResponse"
            output_name = f"{Path(uploaded.name).stem}_redacted.docx"
            return FileResponse(
                open(output_path, "rb"),
                as_attachment=True,
                filename=output_name,
            )

        except Exception as exc:
            tb = traceback.format_exc()
            logger.error("Redaction failed at step [%s]: %s\n%s", step, exc, tb)
            # Show the real error so we can diagnose it
            error_msg = (
                f"Failed at step: <strong>{step}</strong><br><br>"
                f"Error type: <code>{type(exc).__name__}</code><br>"
                f"Error detail: <code>{exc}</code>"
            )
            return render(
                request,
                "redactor/index.html",
                {"error": error_msg},
            )

    return render(request, "redactor/index.html")


def debug(request):
    """
    Diagnostic endpoint — visit /debug/ to see system status.
    Shows Python/Django/spaCy versions and whether imports work.
    """
    import sys
    import django
    lines = []

    lines.append(f"Python: {sys.version}")
    lines.append(f"Django: {django.__version__}")
    lines.append(f"Temp dir: {tempfile.gettempdir()}")
    lines.append("")

    # Test redact_pii import
    try:
        import redact_pii
        lines.append("redact_pii import: OK")
    except Exception as e:
        lines.append(f"redact_pii import: FAILED — {e}")
        lines.append(traceback.format_exc())

    # Test spaCy
    try:
        import spacy
        lines.append(f"spaCy version: {spacy.__version__}")
        nlp = spacy.load("en_core_web_sm", exclude=["tagger", "parser", "lemmatizer", "attribute_ruler"])
        lines.append("spaCy model load: OK")
        doc = nlp("John Smith lives in Mumbai.")
        ents = [(e.text, e.label_) for e in doc.ents]
        lines.append(f"spaCy NER test: {ents}")
    except Exception as e:
        lines.append(f"spaCy: FAILED — {e}")
        lines.append(traceback.format_exc())

    # Test python-docx
    try:
        from docx import Document
        lines.append("python-docx import: OK")
    except Exception as e:
        lines.append(f"python-docx: FAILED — {e}")

    # Test whitenoise
    try:
        import whitenoise
        lines.append(f"whitenoise: OK ({whitenoise.__version__})")
    except Exception as e:
        lines.append(f"whitenoise: FAILED — {e}")

    # Test Faker
    try:
        from faker import Faker
        lines.append("Faker import: OK")
    except Exception as e:
        lines.append(f"Faker: FAILED — {e}")

    return HttpResponse("<pre>" + "\n".join(lines) + "</pre>", content_type="text/html")


def _redacted_document(request):
    """Return the current visitor's generated document, if it still exists."""
    path = request.session.get("redacted_document_path")

    if not path or not Path(path).is_file():
        return None

    return path


def preview(request):
    output_path = _redacted_document(request)
    if not output_path:
        return redirect("index")

    doc = Document(output_path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            paragraphs.extend(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )

    return render(
        request,
        "redactor/preview.html",
        {"paragraphs": paragraphs},
    )


def download(request):
    output_path = _redacted_document(request)
    if not output_path:
        return redirect("index")

    return FileResponse(
        open(output_path, "rb"),
        as_attachment=True,
        filename=request.session.get(
            "redacted_document_name", "redacted_document.docx"
        ),
    )
