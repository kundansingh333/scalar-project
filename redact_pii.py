import os
import re
import shutil
import tempfile
import zipfile

import spacy
from faker import Faker
from PIL import Image
from docx import Document

# =====================================================
# Configuration
# =====================================================

nlp = spacy.load("en_core_web_sm")
nlp.max_length = 5_000_000

fake = Faker()

# Common document/legal terms that should not be replaced
ORG_ALLOWLIST = {
    "RED HERRING",
    "RED HERRING PROSPECTUS",
    "PROSPECTUS",
    "Registrar of Companies",
    "Registrar of Companies Maharashtra",
    "Board",
    "Corporate Office",
    "CURRENCY",
    "Tower 2",
    "Village Birdewadi",
    "Maharashtra",
    "Pune",
}

# =====================================================
# Mapping Dictionaries (Consistency)
# =====================================================

name_map = {}
email_map = {}
phone_map = {}
company_map = {}
location_map = {}
date_map = {}

# =====================================================
# Fake Value Generators
# =====================================================

def fake_name(real_name):
    if real_name not in name_map:
        name_map[real_name] = fake.name()
    return name_map[real_name]


def fake_email(real_email):
    if real_email not in email_map:
        email_map[real_email] = fake.email()
    return email_map[real_email]


def fake_phone(real_phone):
    if real_phone not in phone_map:
        phone_map[real_phone] = "+91 " + "".join(
            str(fake.random_digit()) for _ in range(10)
        )
    return phone_map[real_phone]


def fake_company(real_company):
    if real_company in ORG_ALLOWLIST:
        return real_company

    if real_company not in company_map:
        company_map[real_company] = fake.company()

    return company_map[real_company]


def fake_location(real_location):
    if real_location not in location_map:
        location_map[real_location] = (
            fake.city() + ", " + fake.state()
        )
    return location_map[real_location]


def fake_date(real_date):
    if real_date not in date_map:
        date_map[real_date] = fake.date(pattern="%B %d, %Y")
    return date_map[real_date]

# =====================================================
# Regex Patterns
# =====================================================

EMAIL_PATTERN = re.compile(
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
)

PHONE_PATTERN = re.compile(
    r"\+?\s?91[\s-]?\d{2}[\s-]?\d{4}[\s-]?\d{4}"
    r"|\+?\s?91[\s-]?\d{2}[\s-]?\d{8}"
    r"|\b\d{10}\b"
)

IP_PATTERN = re.compile(
    r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
)

DOB_PATTERN = re.compile(
    r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b"
)

SSN_PATTERN = re.compile(
    r"\b\d{3}-\d{2}-\d{4}\b"
)

CREDIT_CARD_PATTERN = re.compile(
    r"\b(?:\d[ -]*?){13,16}\b"
)

# =====================================================
# Entity Cleaning
# =====================================================

INVALID_PERSON_WORDS = {
    "Taluka",
    "Village",
    "Road",
    "Street",
    "Centre",
    "Office",
    "Tower",
    "Business",
    "Pune",
    "Mumbai",
    "Baner",
    "Khed",
    "Maharashtra",
}

INVALID_ORGS = {
    "Anchor Investors",
    "Bid/Offer Closing Day",
    "Board",
    "Corporate Office",
    "CURRENCY",
}

def clean_persons(persons):
    cleaned = []

    for p in persons:
        if any(word in p for word in INVALID_PERSON_WORDS):
            continue

        if len(p.split()) >= 2:
            cleaned.append(p)

    return list(dict.fromkeys(cleaned))


def clean_orgs(orgs):
    cleaned = []

    for o in orgs:
        if o in INVALID_ORGS:
            continue

        cleaned.append(o)

    return list(dict.fromkeys(cleaned))

# =====================================================
# Text Normalization
# =====================================================

def normalize_text(text):
    text = re.sub(
        r"(?<=[A-Za-z])\n(?=[A-Za-z])",
        "",
        text
    )

    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)

    return text.strip()

# =====================================================
# Document Loading
# =====================================================

def load_document(path):
    return Document(path)

# =====================================================
# Regex Detection
# =====================================================

def detect_regex_pii(text):
    return {
        "emails": EMAIL_PATTERN.findall(text),
        "phones": PHONE_PATTERN.findall(text),
        "ips": IP_PATTERN.findall(text),
        "dobs": DOB_PATTERN.findall(text),
        "ssns": SSN_PATTERN.findall(text),
        "credit_cards": CREDIT_CARD_PATTERN.findall(text),
    }

# =====================================================
# spaCy Detection
# =====================================================

def detect_spacy_entities(text):
    entities = {
        "persons": [],
        "organizations": [],
        "locations": [],
        "dates": [],
    }

    chunk_size = 50_000

    for i in range(0, len(text), chunk_size):
        chunk = text[i:i + chunk_size]
        doc = nlp(chunk)

        for ent in doc.ents:
            if ent.label_ == "PERSON":
                entities["persons"].append(ent.text)

            elif ent.label_ == "ORG":
                entities["organizations"].append(ent.text)

            elif ent.label_ in ("GPE", "LOC", "FAC"):
                entities["locations"].append(ent.text)

            elif ent.label_ == "DATE":
                entities["dates"].append(ent.text)

    entities["persons"] = clean_persons(entities["persons"])
    entities["organizations"] = clean_orgs(
        entities["organizations"]
    )

    return entities

# =====================================================
# Redaction Engine & DOCX Processing
# =====================================================

def redact_document(doc):
    """
    Extracts all text globally, runs expensive NLP once, 
    and applies fast string replacements across all document runs.
    This prevents OOM kills and timeouts on large documents.
    """
    # 1. Gather all text for global context
    all_text = []
    
    def extract_text_from_paragraphs(paragraphs):
        for p in paragraphs:
            if p.text.strip():
                all_text.append(normalize_text(p.text))
                
    extract_text_from_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_text_from_paragraphs(cell.paragraphs)
                
    full_text = "\n".join(all_text)
    
    # 2. Global detection (NLP runs ONCE)
    regex_entities = detect_regex_pii(full_text)
    spacy_entities = {"persons": [], "organizations": [], "locations": [], "dates": []}
    
    if len(full_text) <= 5_000_000:  # spacy max length limit
        spacy_entities = detect_spacy_entities(full_text)
        
    # Pre-sort entities by length (longest first) to prevent partial replacements
    emails = sorted(set(regex_entities["emails"]), key=len, reverse=True)
    phones = sorted(set(regex_entities["phones"]), key=len, reverse=True)
    
    persons = sorted(spacy_entities["persons"], key=len, reverse=True)
    orgs = sorted(spacy_entities["organizations"], key=len, reverse=True)
    locs = sorted(spacy_entities["locations"], key=len, reverse=True)
    dates = []
    for d in spacy_entities["dates"]:
        if re.match(r"^[A-Za-z]+ \d{1,2}, \d{4}$", d):
            dates.append(d)
    dates = sorted(dates, key=len, reverse=True)

    # 3. Fast global replacement function
    def apply_redaction(text):
        if not text.strip():
            return text
            
        text = normalize_text(text)
        
        for email in emails:
            if email in text: text = text.replace(email, fake_email(email))
        for phone in phones:
            if phone in text: text = text.replace(phone, fake_phone(phone))
            
        for person in persons:
            if person in text: text = text.replace(person, fake_name(person))
        for org in orgs:
            if org in text: text = text.replace(org, fake_company(org))
        for loc in locs:
            if loc in text: text = text.replace(loc, fake_location(loc))
        for date in dates:
            if date in text: text = text.replace(date, fake_date(date))
            
        return text

    # 4. Apply to all runs
    def redact_paragraphs_in_place(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                if run.text.strip():
                    run.text = apply_redaction(run.text)

    redact_paragraphs_in_place(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                redact_paragraphs_in_place(cell.paragraphs)

    return doc

# =====================================================
# Image Redaction (Only Aadhaar / PAN Images)
# =====================================================

def blackout_image(image_path):
    """Replace the image with a solid black image of the same size."""
    image = Image.open(image_path)
    black = Image.new(image.mode, image.size, color="black")
    black.save(image_path)


def blackout_identity_images_in_docx(docx_path):
    """
    Black out only the PAN and Aadhaar images (image4.png and image5.png).
    All logos and other images remain unchanged.
    Uses proper temp directories so concurrent requests don't conflict.
    """
    temp_dir = tempfile.mkdtemp(prefix="pii_redact_")

    try:
        with zipfile.ZipFile(docx_path, "r") as zip_ref:
            zip_ref.extractall(temp_dir)

        media_dir = os.path.join(temp_dir, "word", "media")
        identity_images = {"image4.png", "image5.png"}

        if os.path.exists(media_dir):
            for filename in os.listdir(media_dir):
                if filename in identity_images:
                    blackout_image(os.path.join(media_dir, filename))

        # Rebuild the DOCX into a fresh temp file then replace original
        with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
            temp_zip = tmp.name

        with zipfile.ZipFile(temp_zip, "w", zipfile.ZIP_DEFLATED) as zip_out:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    arcname = os.path.relpath(full_path, temp_dir)
                    zip_out.write(full_path, arcname)

        os.replace(temp_zip, docx_path)

    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


# =====================================================
# Main (standalone script usage)
# =====================================================
def main():
    input_path = "input/Red Herring Prospectus.docx"
    output_path = "output/Red_Herring_Prospectus_Redacted.docx"

    doc = load_document(input_path)
    redacted_doc = redact_document(doc)
    redacted_doc.save(output_path)

    blackout_identity_images_in_docx(output_path)

    print("Redaction complete!")
    print(f"Output saved to: {output_path}")


if __name__ == "__main__":
    main()
