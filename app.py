import logging
import os
import tempfile
import traceback
import uuid
from pathlib import Path

from docx import Document
from flask import (
    Flask,
    abort,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)

from redact_pii import blackout_identity_images_in_docx, redact_document

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "change-me-in-production")
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB upload limit

logging.basicConfig(
    level=logging.INFO,
    format="[%(levelname)s] %(name)s: %(message)s",
)
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# In-process file store
# Single-worker deployment (workers=1 in gunicorn.conf.py) means this dict
# persists across requests for the lifetime of the process.
# ---------------------------------------------------------------------------

_store: dict[str, dict] = {}


def _extract_preview(docx_path: str, max_paragraphs: int = 120) -> list[str]:
    """Return plain-text paragraphs from a DOCX for browser preview."""
    doc = Document(docx_path)
    items: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            items.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    items.append(cell.text.strip())

    return items[:max_paragraphs]


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/redact", methods=["POST"])
def redact():
    uploaded = request.files.get("document")
    if not uploaded or not uploaded.filename:
        return render_template("index.html", error="Please select a .docx file to upload.")

    if not uploaded.filename.lower().endswith(".docx"):
        return render_template("index.html", error="Only .docx files are supported.")

    step = "initializing"
    input_path = None

    try:
        step = "saving upload to temp file"
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
            uploaded.save(tmp_in)
            input_path = tmp_in.name

        output_path = str(
            Path(tempfile.gettempdir()) / f"redacted-{uuid.uuid4()}.docx"
        )

        step = "opening document"
        doc = Document(input_path)
        log.info("Opened: %s", uploaded.filename)

        step = "redacting text (NLP + regex)"
        redacted = redact_document(doc)
        redacted.save(output_path)
        log.info("Redaction complete → %s", output_path)

        step = "blacking out identity images"
        blackout_identity_images_in_docx(output_path)
        log.info("Image blackout done")

        step = "building preview"
        preview = _extract_preview(output_path)

        file_id = str(uuid.uuid4())
        _store[file_id] = {
            "path": output_path,
            "name": f"{Path(uploaded.filename).stem}_redacted.docx",
            "preview": preview,
        }

        return redirect(url_for("result", file_id=file_id))

    except Exception as exc:
        log.error("Failed at [%s]: %s\n%s", step, exc, traceback.format_exc())
        error = (
            f"Redaction failed at step: <strong>{step}</strong><br>"
            f"<code>{type(exc).__name__}: {exc}</code>"
        )
        return render_template("index.html", error=error), 500

    finally:
        if input_path and os.path.exists(input_path):
            try:
                os.remove(input_path)
            except OSError:
                pass


@app.route("/result/<file_id>")
def result(file_id):
    info = _store.get(file_id)
    if not info or not Path(info["path"]).is_file():
        return redirect(url_for("index"))

    return render_template(
        "result.html",
        file_id=file_id,
        filename=info["name"],
        paragraphs=info["preview"],
    )


@app.route("/download/<file_id>")
def download(file_id):
    info = _store.get(file_id)
    if not info or not Path(info["path"]).is_file():
        abort(404)

    return send_file(
        info["path"],
        as_attachment=True,
        download_name=info["name"],
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---------------------------------------------------------------------------
# Entry point (dev only — production uses gunicorn)
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
