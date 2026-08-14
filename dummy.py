from docx import Document
import redact_pii

doc = Document()
doc.add_paragraph("My name is John Doe and my email is john@example.com.")
doc.save("test_input.docx")

try:
    redact_pii.redact_document(Document("test_input.docx"))
    print("Success!")
except Exception as e:
    import traceback
    traceback.print_exc()
